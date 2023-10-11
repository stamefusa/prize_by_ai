from docx import Document
import subprocess
import os

list = [
    "デイリーポータルＺライター賞",
    "爲房新太朗様",
    "あなたの深く考えられた記事と、緻密で愛情ある文章に感銘を受け、デイリーポータルＺライター賞を授与いたします。\nあなたの寄稿は常に読者の注目を集め、彼らに新たな視点を提供しています。\n更なるご成功を心よりお祈りしております。",
    "授与日：2023年7月10日",
    "授与組織：日本メディア協会"
]

tmp = list[2].split("\n")
count = 0 # この文章がWORDファイル上で何行になるかを示す数
for t in tmp:
    print(t)
    count += int(len(t)/28) # このWORDファイルの一行の最大文字数は27文字。一文が長く文中で改行される回数を足す
    count += 1 # 句点ごとに改行になるのでその分足す
print(count)
lf_count = 11 - count

doc = Document("template.docx")
print(doc.paragraphs[2].text)
print(doc.paragraphs[2].runs[0].text)
doc.paragraphs[0].runs[0].text = list[0]
doc.paragraphs[4].runs[0].text = list[2]
doc.paragraphs[6].runs[0].text = "\n" * lf_count + list[3]
#if len(list[2]) > 100 and len(list[2]) < 200:

doc.save("sample.docx")
filepath = os.environ['HOME'] + "/prize_by_ai/sample.docx"
subprocess.run("docx2pdf " + filepath, shell=True)

# doc.add_paragraph(textwrap.dedent('''デイリーポータルＺライター賞

# 爲房新太朗様

# あなたの深く考えられた記事と、緻密で愛情ある文章に感銘を受け、デイリーポータルＺライター賞を授与いたします。あなたの寄稿は常に読者の注目を集め、彼らに新たな視点を提供しています。更なるご成功を心よりお祈りしております。

# 授与日：2021年7月10日
# 授与組織：日本メディア協会'''))
#doc.add_paragraph(list[0])
#doc.save("sample.docx")