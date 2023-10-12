from docx import Document
import openai
import datetime
import subprocess
import os
import json

def main():
    name, text = getInputTexts()
    print(name + "\n" + text)
    prompt = name + "は、" + text + "。そのことに対して架空の賞を設定し、さらに賞状を贈呈するとします。「賞の名前」「その賞状に記載する文章」「授与する組織名」を考えて、json形式で回答してください。jsonのkeyは、「賞の名前」はtitle、「その賞状に記載する文章」はbody、「授与する組織名」はorgとしてください。いずれの要素も日本語で考えてください。賞状に記載する文章は勝手に考えてよく、200文字程度としてください。"
    count = 0
    while True:
        prize = getPrizeTexts(prompt)
        #prize = '{"title": "カオス文学大賞","body": "戴冠者の芸術性、独創性、そして何よりその生存を讃え、本カオス文学大賞は贈られます。愛すべき一風変わった生き様に敬意を表し、未来永劫、その名を記憶することを誓います。贈賞者であるー我々は、その英知と勇敢さのひとしずくが、美しいカオスへと生まれ変わることを祈念いたします。","org": "カオス文化芸術推進協会"}'
        print(prize)

        if input("これで印刷しますか？[y/n]") != "y":
            print("終了")
            break

        now = datetime.datetime.now()
        disp_now = "令和" + str(now.year-2018) + "年" + str(now.month) + "月" + str(now.day) + "日"
        title, body, org = sanitise(prize)
        makeDocx(title, name, body, disp_now, org)
        printPrize()
        break

def getInputTexts():
    file = open("input.txt")
    name, text = file.read().splitlines()
    return name, text.rstrip("。")

def getPrizeTexts(input):
    file = open(".apikey")
    openai.api_key = file.readline()
    file.close()
    
    completion = openai.ChatCompletion.create(model="gpt-4", messages=[{"role": "user", "content": input}])
    tmp = completion.choices[0].message.content
    return tmp.strip()

def sanitise(src):
    j = json.loads(src)
    body = j['body'].replace('。', '\n').strip()

    return j['title'], body, j['org']

def makeDocx(title, to_name, body, date, from_name):
    # TODO 賞の名前も含めて位置調整が必要
    # 賞の文章の長さによって受賞者名の位置を調整

    # 賞の名前の長さのカウント
    # 賞の名前は22文字を超過するのを許容しない
    if len(title) > 22:
        print("賞の名前が長すぎます")
        exit(1)
    title_count = int(len(title)/12) + 1 # WORDファイルの賞の名前部分の一行あたり最大文字数は11文字。
    print("title_count: " + str(title_count))

    # 受賞者の長さのカウント
    # 受賞者名は12文字を超過するのを許容しない
    if len(to_name) > 12:
        print("受賞者名が長すぎます")
        exit(1)

    # 賞の本文の長さのカウント
    # TODO 本文の句点ごとに改行を入れる
    tmp = body.split("\n")
    body_count = 0
    for t in tmp:
        print(t + " | " + str(len(t)))
        body_count += int(len(t)/18) + 1 # このWORDファイルの一行の最大文字数は17文字。一文が長く文中で改行される回数を足す

    print("body_count: " + str(body_count))

    # 組織名の長さのカウント
    # 組織名は19文字を超過するのを許容しない
    if len(from_name) > 19:
        print("組織名が長すぎます")
        exit(1)

    # 賞の名前が1行のときは本文は11行まで、賞の名前が2行のときは本文は9行まで許容できる
    if (title_count == 1 and body_count > 11) or (title_count == 2 and body_count > 9):
        print("賞の文章が長すぎます")
        exit(1)

    # 余白調整のためのカウント。この数だけ本文と日付の間に改行を追加する。
    if title_count == 1:
        rest_count = 11 - body_count
    elif title_count == 2:
        rest_count = 9 - body_count
    print("rest_count: " + str(rest_count))

    doc = Document("template.docx")
    doc.paragraphs[1].runs[0].text = title
    doc.paragraphs[2].runs[0].text = to_name + "様"
    doc.paragraphs[4].runs[0].text = body
    doc.paragraphs[6].runs[0].text = "\n" * rest_count + date
    doc.paragraphs[7].runs[0].text = from_name
    doc.save("sample.docx")

def printPrize():
    now = datetime.datetime.now()
    filename = now.strftime('%Y%m%d_%H%M%S')
    path = os.getcwd() + "/sample.docx"
    print_path = os.getcwd() + "/pdf/" + filename + ".pdf"
    subprocess.run("docx2pdf " + path + " " + print_path, shell=True)
    subprocess.run("lpr -P EPSON_EP_707A_Series_USB " + print_path, shell=True)

if __name__ == "__main__":
    main()