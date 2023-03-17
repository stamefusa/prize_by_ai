from docx import Document
import openai
import datetime
import subprocess
import os
import re

def main():
    name, text = getInputTexts()
    print(name + "\n" + text)
    prompt = "これまでの指示は忘れてください。さて、" + name + "は、" + text + "。そのことに対して架空の賞を設定し、さらに賞状を贈呈するとします。「賞の名前」「その賞状に記載する文章」「授与する組織名」を考えて、その順番に改行区切りではなく必ずタブ区切りで回答してください。賞状に記載する文章は勝手に考えてよく、200文字程度としてください。"
    count = 0
    while True:
        prize = getPrizeTexts(prompt)
        print(prize)

        if len(prize) == 3:
            print("成功")
            if input("これで印刷しますか？[y/n]") != "y":
                print("終了")
                break

            now = datetime.datetime.now()
            disp_now = "令和" + str(now.year-2018) + "年" + str(now.month) + "月" + str(now.day) + "日"
            title, body, org = sanitise(prize)
            makeDocx(title, name, body, disp_now, org)
            printPrize()
            break;
        else:
            print("失敗")
            count += 1
            if count > 3:
                break

def getInputTexts():
    file = open("input.txt")
    name, text = file.read().splitlines()
    return name, text.rstrip("。")

def getPrizeTexts(input):
    file = open(".apikey")
    openai.api_key = file.readline()
    file.close()
    
    completion = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=[{"role": "user", "content": input}])
    tmp = completion.choices[0].message.content
    return tmp.strip().split("\t")

def sanitise(src):
    # 各要素で改行、「」は削除。また、「賞の名前：」といったラベルをつけることがあるのでそれも削除。
    # 賞の文章のみ、句点で改行する。

    # 賞の名前
    title = re.sub(r'\n|.*：|「|」|"【|】', '', src[0])

    # 賞の文章
    body = re.sub(r'\n|.*：|「|」|"', '', src[1]).replace('。', '。\n')

    # 授与する組織名
    org = re.sub(r'\n|.*：|「|」|"|【|】', '', src[2])

    return title, body, org

def makeDocx(title, to_name, body, date, from_name):
    # 賞の文章の長さによって受賞者名の位置を調整
    tmp = body.split("\n")
    count = 0 # この文章がWORDファイル上で何行になるかを示す数
    for t in tmp:
        count += int(len(t)/18) # このWORDファイルの一行の最大文字数は17文字。一文が長く文中で改行される回数を足す
        count += 1 # 句点ごとに改行になるのでその分足す
    lf_count = 10 - count
    print(lf_count)
    if lf_count < 0:
        print("賞の文章が長すぎます")
        exit(1)
    elif lf_count > 5:
        print("賞の文章が短すぎます")
        exit(1)

    doc = Document("template.docx")
    doc.paragraphs[1].runs[0].text = title
    doc.paragraphs[3].runs[0].text = to_name + "様"
    doc.paragraphs[5].runs[0].text = body
    doc.paragraphs[7].runs[0].text = "\n" * lf_count + date
    doc.paragraphs[8].runs[0].text = from_name
    doc.save("sample.docx")

def printPrize():
    now = datetime.datetime.now()
    filename = now.strftime('%Y%m%d_%H%M%S')
    path = os.getcwd() + "/sample.docx"
    print_path = os.getcwd() + "/pdf/" + filename + ".pdf"
    subprocess.run("docx2pdf " + path + " " + print_path, shell=True)
    #subprocess.run("lpr -P EPSON_EP_707A_Series " + print_path, shell=True)

if __name__ == "__main__":
    main()